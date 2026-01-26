"""
Inline element rendering for markdown to DOCX conversion.
"""

from docx.shared import Pt


class InlineRenderer:
    """Renders inline markdown elements to DOCX runs."""
    
    def __init__(self, builder):
        self.builder = builder
    
    def render_children(self, children: list, paragraph):
        """Render list of child nodes."""
        for child in children:
            self.render_inline(child, paragraph)
    
    def render_inline(self, node: dict, paragraph, bold: bool = False, italic: bool = False):
        """Render inline node."""
        ntype = node.get("type")
        
        if ntype == "text":
            self._render_text(node, paragraph, bold, italic)
        elif ntype == "strong":
            for c in node.get("children", []):
                self.render_inline(c, paragraph, bold=True, italic=italic)
        elif ntype == "emphasis":
            for c in node.get("children", []):
                self.render_inline(c, paragraph, bold=bold, italic=True)
        elif ntype == "codespan":
            self._render_codespan(node, paragraph)
        elif ntype == "link":
            self._render_link(node, paragraph, bold, italic)
        elif ntype in ["softbreak", "linebreak"]:
            paragraph.add_run("\n")
    
    def _render_text(self, node: dict, paragraph, bold: bool, italic: bool):
        """Render text node with formula support."""
        text = node.get("raw", "")
        
        if self.builder.latex.block_marker_pattern.search(text):
            self.builder.process_block_markers(text, paragraph)
        elif self.builder.latex.inline_marker_pattern.search(text):
            self.builder.process_inline_markers(text, paragraph, bold=bold, italic=italic)
        else:
            self.builder.add_text_run(paragraph, text, bold=bold, italic=italic)
    
    def _render_codespan(self, node: dict, paragraph):
        """Render inline code."""
        run = paragraph.add_run(node.get("raw", ""))
        run.font.name = "Courier New"
        run.font.size = Pt(12)
    
    def _render_link(self, node: dict, paragraph, bold: bool, italic: bool):
        """Render link."""
        text = "".join(
            c.get("raw", "") for c in node.get("children", []) 
            if c.get("type") == "text"
        )
        url = node.get("attrs", {}).get("url", "")
        run = paragraph.add_run(f"{text} ({url})")
        run.bold = bold
        run.italic = italic
    
    def render_children_in_cell(self, children: list, paragraph, max_width: float):
        """Render children in a table cell with width constraint."""
        for child in children:
            self.render_inline_in_cell(child, paragraph, max_width)
    
    def render_inline_in_cell(self, node: dict, paragraph, max_width: float, 
                               bold: bool = False, italic: bool = False):
        """Render inline node in a table cell with width constraint."""
        ntype = node.get("type")
        
        if ntype == "text":
            self._render_text_in_cell(node, paragraph, max_width, bold, italic)
        elif ntype == "strong":
            for c in node.get("children", []):
                self.render_inline_in_cell(c, paragraph, max_width, bold=True, italic=italic)
        elif ntype == "emphasis":
            for c in node.get("children", []):
                self.render_inline_in_cell(c, paragraph, max_width, bold=bold, italic=True)
        elif ntype == "codespan":
            self._render_codespan(node, paragraph)
        elif ntype == "link":
            self._render_link(node, paragraph, bold, italic)
    
    def _render_text_in_cell(self, node: dict, paragraph, max_width: float, 
                              bold: bool, italic: bool):
        """Render text node in table cell."""
        text = node.get("raw", "")
        
        if self.builder.latex.block_marker_pattern.search(text):
            self.builder.process_block_markers(text, paragraph)
        elif self.builder.latex.inline_marker_pattern.search(text):
            self.builder.process_inline_markers(text, paragraph, bold=bold, 
                                                  italic=italic, max_width=max_width)
        elif self.builder.latex.inline_pattern.search(text):
            self._render_text_with_inline_formulas(text, paragraph, bold, italic, max_width)
        else:
            self.builder.add_text_run(paragraph, text, bold=bold, italic=italic)
    
    def _render_text_with_inline_formulas(self, text: str, paragraph, 
                                           bold: bool, italic: bool, max_width: float):
        """Render text containing raw $...$ inline formulas."""
        pattern = self.builder.latex.inline_pattern
        last_end = 0
        
        for match in pattern.finditer(text):
            if match.start() > last_end:
                before_text = text[last_end:match.start()]
                self.builder.add_text_run(paragraph, before_text, bold=bold, italic=italic)
            
            formula = match.group(1)
            self.builder.formulas.add_inline_formula(formula, paragraph, bold=bold, 
                                                      italic=italic, max_width=max_width)
            last_end = match.end()
        
        if last_end < len(text):
            self.builder.add_text_run(paragraph, text[last_end:], bold=bold, italic=italic)
