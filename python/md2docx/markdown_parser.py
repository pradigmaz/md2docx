"""
Markdown parsing and AST processing module.
"""

import re
import mistune
from .document_builder import DocumentBuilder
from docx.shared import Pt, Cm


class MarkdownProcessor:
    """Processes markdown AST and builds DOCX document."""
    
    def __init__(self, builder: DocumentBuilder):
        self.builder = builder
    
    def _fix_broken_table_lines(self, text: str) -> str:
        """Fix table lines that were broken by text wrapping.
        
        Some editors wrap long lines, breaking markdown tables.
        This joins lines that appear to be continuations of table rows.
        """
        lines = text.split('\n')
        fixed_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # Check if this looks like a table row (starts with |)
            if line.strip().startswith('|'):
                # A complete table row should end with |
                # Keep joining lines until we get a complete row
                while not line.rstrip().endswith('|') and i + 1 < len(lines):
                    next_line = lines[i + 1]
                    # If next line starts with |, it's a new row, stop joining
                    if next_line.strip().startswith('|'):
                        break
                    # If next line is empty, stop joining
                    if not next_line.strip():
                        break
                    # Join the continuation
                    line = line.rstrip() + next_line.lstrip()
                    i += 1
            
            fixed_lines.append(line)
            i += 1
        
        return '\n'.join(fixed_lines)
    
    def parse(self, text: str) -> list:
        """Parse markdown text to AST."""
        # Fix broken table lines first
        text = self._fix_broken_table_lines(text)
        
        # Extract LaTeX before parsing (prevents mistune from breaking formulas)
        text, self.builder.latex_blocks = self.builder.latex.extract_blocks(text)
        text, self.builder.latex_inlines = self.builder.latex.extract_inlines(text)
        
        md = mistune.create_markdown(renderer=None, plugins=['table'])
        return md(text)
    
    def process(self, ast: list):
        """Process AST nodes."""
        self._prev_list_type = None
        for node in ast:
            self._process_node(node)
            # Track if this was a list for continuation
            if node.get("type") == "list":
                ordered = node.get("attrs", {}).get("ordered", False)
                self._prev_list_type = "ordered" if ordered else "bullet"
            elif node.get("type") not in ["blank_line"]:
                # Reset if non-list, non-blank node
                if node.get("type") != "paragraph" or not self._is_block_formula_only(node):
                    self._prev_list_type = None
    
    def _is_block_formula_only(self, node: dict) -> bool:
        """Check if paragraph contains only a block formula marker."""
        children = node.get("children", [])
        if len(children) == 1 and children[0].get("type") == "text":
            text = children[0].get("raw", "").strip()
            return self.builder.latex.block_marker_pattern.fullmatch(text) is not None
        return False
    
    def _process_node(self, node: dict, list_type: str = None, is_first: bool = False):
        """Process single AST node."""
        ntype = node.get("type")
        
        if ntype == "heading":
            self._handle_heading(node)
        elif ntype == "paragraph":
            self._handle_paragraph(node)
        elif ntype == "list":
            self._handle_list(node)
        elif ntype == "list_item":
            self._handle_list_item(node, list_type, is_first)
        elif ntype == "block_code":
            self._handle_code_block(node)
        elif ntype == "block_quote":
            self._handle_blockquote(node)
        elif ntype == "thematic_break":
            pass  # Ignore --- separators
        elif ntype == "blank_line":
            pass  # Ignore blank lines - spacing handled by paragraph format
        elif ntype == "table":
            self._handle_table(node)
    
    def _handle_heading(self, node: dict):
        """Handle heading node."""
        from docx.shared import RGBColor
        
        level = node.get("attrs", {}).get("level", 1)
        p = self.builder.add_heading("", level)
        self._render_children(node.get("children", []), p)
        
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
    
    def _handle_paragraph(self, node: dict):
        """Handle paragraph node."""
        children = node.get("children", [])
        
        # Filter out empty breaks at start/end
        while children and children[0].get("type") in ["softbreak", "linebreak"]:
            children = children[1:]
        while children and children[-1].get("type") in ["softbreak", "linebreak"]:
            children = children[:-1]
        
        if not children:
            return  # Skip empty paragraphs
        
        # Check if paragraph contains only block formula markers
        if len(children) == 1 and children[0].get("type") == "text":
            text = children[0].get("raw", "").strip()
            if self.builder.latex.block_marker_pattern.fullmatch(text):
                # Just render the block formula without creating a paragraph
                self.builder.process_block_markers(text, None)
                return
        
        # Check if paragraph contains block formula mixed with text
        if len(children) == 1 and children[0].get("type") == "text":
            text = children[0].get("raw", "")
            if self.builder.latex.block_marker_pattern.search(text):
                # Split text around block markers and process separately
                self._handle_mixed_block_paragraph(text)
                return
        
        # Split by softbreak/linebreak - each line becomes a new paragraph
        current_children = []
        for child in children:
            ctype = child.get("type")
            if ctype in ["softbreak", "linebreak"]:
                # Render current children as paragraph, then start new one
                if current_children:
                    p = self.builder.add_paragraph()
                    for c in current_children:
                        self._render_inline(c, p)
                    current_children = []
            else:
                current_children.append(child)
        
        # Render remaining children
        if current_children:
            p = self.builder.add_paragraph()
            for c in current_children:
                self._render_inline(c, p)
    
    def _handle_mixed_block_paragraph(self, text: str):
        """Handle paragraph with mixed text and block formulas."""
        import re
        parts = []
        last_end = 0
        
        for m in self.builder.latex.block_marker_pattern.finditer(text):
            if m.start() > last_end:
                parts.append(("text", text[last_end:m.start()]))
            parts.append(("block", int(m.group(1))))
            last_end = m.end()
        
        if last_end < len(text):
            parts.append(("text", text[last_end:]))
        
        for ptype, content in parts:
            if ptype == "text":
                content = content.strip()
                if content:
                    p = self.builder.add_paragraph()
                    if self.builder.latex.inline_marker_pattern.search(content):
                        self.builder.process_inline_markers(content, p)
                    else:
                        self.builder.add_text_run(p, content)
            elif ptype == "block":
                formula = self.builder.formulas.latex_blocks.get(content, "")
                self.builder.formulas.add_block(formula)
    
    def _handle_list(self, node: dict):
        """Handle list node."""
        ordered = node.get("attrs", {}).get("ordered", False)
        list_type = "List Number" if ordered else "List Bullet"
        current_type = "ordered" if ordered else "bullet"
        
        children = node.get("children", [])
        
        for idx, item in enumerate(children):
            # Restart only for first item of first list (not continuation)
            restart = (idx == 0 and ordered and self._prev_list_type != current_type)
            self._process_node(item, list_type=list_type, is_first=restart)
    
    def _handle_list_item(self, node: dict, list_type: str, is_first: bool = False):
        """Handle list item node."""
        ordered = list_type == "List Number"
        p = self.builder.add_list_item(ordered=ordered, restart=(is_first and ordered))
        
        for child in node.get("children", []):
            ctype = child.get("type")
            if ctype in ["paragraph", "block_text"]:
                self._render_children(child.get("children", []), p)
            else:
                self._render_children([child], p)
    
    def _handle_code_block(self, node: dict):
        """Handle code block node."""
        code = node.get("raw", "")
        lang = node.get("attrs", {}).get("info", "")
        style = node.get("style", "")
        
        # Indented text without language is often just continuation text, not code
        if style == "indent" and not lang:
            # Treat as regular paragraphs with formulas
            lines = code.strip().split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    p = self.builder.add_paragraph()
                    if self.builder.latex.block_marker_pattern.search(line):
                        self.builder.process_block_markers(line, p)
                    elif self.builder.latex.inline_marker_pattern.search(line):
                        self.builder.process_inline_markers(line, p)
                    else:
                        self.builder.add_text_run(p, line)
        else:
            self.builder.add_code_block(code, lang)
    
    def _handle_blockquote(self, node: dict):
        """Handle blockquote node."""
        for child in node.get("children", []):
            p = self.builder.add_blockquote()
            if child.get("type") == "paragraph":
                self._render_children(child.get("children", []), p)
            else:
                self._render_children([child], p)
    
    def _handle_table(self, node: dict):
        """Handle table node."""
        head = body = None
        for child in node.get("children", []):
            if child.get("type") == "table_head":
                head = child
            elif child.get("type") == "table_body":
                body = child
        
        if not head:
            return
        
        head_cells = head.get("children", [])
        num_cols = len(head_cells)
        body_rows = body.get("children", []) if body else []
        num_rows = 1 + len(body_rows)
        
        table = self.builder.add_table(num_rows, num_cols)
        
        # Column widths in points for formula scaling (1cm = 28.35pt)
        # For 3-column tables: 1.5cm, 9cm, 6cm
        if num_cols == 3:
            col_widths_pt = [1.5 * 28.35, 9 * 28.35, 6 * 28.35]
        else:
            col_widths_pt = [16.5 / num_cols * 28.35] * num_cols
        
        # Header row
        for col, cell_node in enumerate(head_cells):
            cell = table.rows[0].cells[col]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            # Use smaller max_width for table cells (leave some padding)
            max_w = col_widths_pt[col] - 10 if col < len(col_widths_pt) else 150
            self._render_children_in_cell(cell_node.get("children", []), p, max_w)
            for run in p.runs:
                run.bold = True
        
        # Body rows
        for row_idx, row_node in enumerate(body_rows):
            for col, cell_node in enumerate(row_node.get("children", [])):
                if col < num_cols:
                    cell = table.rows[row_idx + 1].cells[col]
                    p = cell.paragraphs[0]
                    p.paragraph_format.first_line_indent = Cm(0)
                    max_w = col_widths_pt[col] - 10 if col < len(col_widths_pt) else 150
                    self._render_children_in_cell(cell_node.get("children", []), p, max_w)
    
    def _render_children_in_cell(self, children: list, paragraph, max_width: float):
        """Render children in a table cell with width constraint."""
        for child in children:
            self._render_inline_in_cell(child, paragraph, max_width)
    
    def _render_inline_in_cell(self, node: dict, paragraph, max_width: float, bold: bool = False, italic: bool = False):
        """Render inline node in a table cell with width constraint."""
        ntype = node.get("type")
        
        if ntype == "text":
            text = node.get("raw", "")
            
            # Check for markers first (from pre-extracted formulas)
            if self.builder.latex.block_marker_pattern.search(text):
                self.builder.process_block_markers(text, paragraph)
            elif self.builder.latex.inline_marker_pattern.search(text):
                self.builder.process_inline_markers(text, paragraph, bold=bold, italic=italic, max_width=max_width)
            # Check for raw $...$ formulas (not extracted in table cells)
            elif self.builder.latex.inline_pattern.search(text):
                self._render_text_with_inline_formulas(text, paragraph, bold, italic, max_width)
            else:
                self.builder.add_text_run(paragraph, text, bold=bold, italic=italic)
        
        elif ntype == "strong":
            for c in node.get("children", []):
                self._render_inline_in_cell(c, paragraph, max_width, bold=True, italic=italic)
        
        elif ntype == "emphasis":
            for c in node.get("children", []):
                self._render_inline_in_cell(c, paragraph, max_width, bold=bold, italic=True)
        
        elif ntype == "codespan":
            run = paragraph.add_run(node.get("raw", ""))
            run.font.name = "Courier New"
            run.font.size = Pt(12)
        
        elif ntype == "link":
            text = "".join(
                c.get("raw", "") for c in node.get("children", []) 
                if c.get("type") == "text"
            )
            url = node.get("attrs", {}).get("url", "")
            run = paragraph.add_run(f"{text} ({url})")
            run.bold = bold
            run.italic = italic
    
    def _render_text_with_inline_formulas(self, text: str, paragraph, bold: bool, italic: bool, max_width: float):
        """Render text containing raw $...$ inline formulas."""
        pattern = self.builder.latex.inline_pattern
        last_end = 0
        
        for match in pattern.finditer(text):
            # Add text before formula
            if match.start() > last_end:
                before_text = text[last_end:match.start()]
                self.builder.add_text_run(paragraph, before_text, bold=bold, italic=italic)
            
            # Render formula
            formula = match.group(1)
            self.builder.formulas.add_inline_formula(formula, paragraph, bold=bold, italic=italic, max_width=max_width)
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            self.builder.add_text_run(paragraph, text[last_end:], bold=bold, italic=italic)
    
    def _render_children(self, children: list, paragraph):
        """Render list of child nodes."""
        for child in children:
            self._render_inline(child, paragraph)
    
    def _render_inline(self, node: dict, paragraph, bold: bool = False, italic: bool = False):
        """Render inline node."""
        ntype = node.get("type")
        
        if ntype == "text":
            text = node.get("raw", "")
            
            if self.builder.latex.block_marker_pattern.search(text):
                self.builder.process_block_markers(text, paragraph)
            elif self.builder.latex.inline_marker_pattern.search(text):
                self.builder.process_inline_markers(text, paragraph, bold=bold, italic=italic)
            else:
                self.builder.add_text_run(paragraph, text, bold=bold, italic=italic)
        
        elif ntype == "strong":
            for c in node.get("children", []):
                self._render_inline(c, paragraph, bold=True, italic=italic)
        
        elif ntype == "emphasis":
            for c in node.get("children", []):
                self._render_inline(c, paragraph, bold=bold, italic=True)
        
        elif ntype == "codespan":
            run = paragraph.add_run(node.get("raw", ""))
            run.font.name = "Courier New"
            run.font.size = Pt(12)
        
        elif ntype == "link":
            text = "".join(
                c.get("raw", "") for c in node.get("children", []) 
                if c.get("type") == "text"
            )
            url = node.get("attrs", {}).get("url", "")
            run = paragraph.add_run(f"{text} ({url})")
            run.bold = bold
            run.italic = italic
        
        elif ntype in ["softbreak", "linebreak"]:
            paragraph.add_run("\n")
