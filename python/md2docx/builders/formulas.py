"""
LaTeX formula handling for DOCX documents.
Uses PNG images rendered via ziamath for LibreOffice compatibility.
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn as oxml_qn
from docx.oxml import OxmlElement

from ..latex_renderer import LaTeXRenderer


class FormulaBuilder:
    """Handles LaTeX formula rendering and insertion into DOCX documents."""
    
    MULTILINE_ENVS = [
        'cases', 'matrix', 'pmatrix', 'bmatrix', 'vmatrix',
        'array', 'align', 'gather', 'split', 'eqnarray'
    ]
    
    def __init__(self, doc, settings: dict):
        self.doc = doc
        self.settings = settings
        self.latex = LaTeXRenderer()
        self.latex_blocks: dict[int, str] = {}
        self.latex_inlines: dict[int, str] = {}
    
    def is_multiline(self, formula: str) -> bool:
        """Check if formula contains multiline constructs."""
        for env in self.MULTILINE_ENVS:
            if f'\\begin{{{env}}}' in formula:
                return True
        if '\\\\' in formula or '\n' in formula:
            return True
        return False
    
    def add_inline(self, paragraph, formula: str, add_text_run_func, max_width: float = None, bold: bool = False, italic: bool = False) -> bool:
        """Add inline LaTeX formula to paragraph as PNG image.
        
        Args:
            paragraph: The paragraph to add the formula to
            formula: LaTeX formula string
            add_text_run_func: Function to add text runs
            max_width: Maximum width in points (for table cells)
            bold: Whether surrounding text is bold (for fallback)
            italic: Whether surrounding text is italic (for fallback)
        """
        is_multi = self.is_multiline(formula)
        png_path = self.latex.render_to_png(formula, is_block=is_multi)
        
        if png_path:
            run = paragraph.add_run()
            try:
                w, h = self.latex.get_image_size(png_path)
                dpi = 192
                w_pt = w * 72 / dpi
                h_pt = h * 72 / dpi
                
                font_size = self.settings.get("fontSize", 14)
                
                if is_multi:
                    max_w = max_width if max_width else 400
                    scale = min(1.0, max_w / w_pt)
                    run.add_picture(png_path, width=Pt(w_pt * scale), height=Pt(h_pt * scale))
                else:
                    # Scale inline formula - allow taller formulas to avoid excessive width
                    target_h = font_size * 1.8  # Allow more height for better proportions
                    max_w = max_width if max_width else 250
                    # Prioritize width constraint over height
                    scale = min(1.0, max_w / w_pt) if w_pt > max_w else min(1.0, target_h / h_pt)
                    run.add_picture(png_path, width=Pt(w_pt * scale), height=Pt(h_pt * scale))
                return True
            except Exception:
                pass
        
        # Fallback to text
        run = paragraph.add_run(f"${formula}$")
        run.font.name = self.settings["fontFamily"]
        run.font.size = Pt(self.settings["fontSize"])
        run.bold = bold
        run.italic = italic
        return False
    
    def add_block(self, formula: str):
        """Add block LaTeX formula as centered paragraph with PNG."""
        png_path = self.latex.render_to_png(formula, is_block=True)
        
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)
        
        if png_path:
            run = p.add_run()
            try:
                w, h = self.latex.get_image_size(png_path)
                dpi = 192
                w_pt = w * 72 / dpi
                h_pt = h * 72 / dpi
                
                max_w = 14 * 28.35
                scale = min(1.0, max_w / w_pt)
                run.add_picture(png_path, width=Pt(w_pt * scale), height=Pt(h_pt * scale))
                return
            except Exception:
                pass
        
        run = p.add_run(f"${formula}$")
        run.font.name = self.settings["fontFamily"]
    
    def process_inline_markers(self, text: str, paragraph, add_text_run_func=None, bold: bool = False, italic: bool = False, max_width: float = None):
        """Process text with ⟦LATEX_INLINE:id⟧ markers.
        
        Args:
            text: Text with inline markers
            paragraph: Paragraph to add content to
            add_text_run_func: Function to add text runs
            bold: Whether text should be bold
            italic: Whether text should be italic
            max_width: Maximum width for formulas in points (for table cells)
        """
        parts = self.latex.inline_marker_pattern.split(text)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part:
                    # Replace regular spaces with non-breaking spaces near brackets
                    # to prevent Word from breaking lines between bracket and formula
                    # \u00A0 is non-breaking space
                    part = part.replace(' (', '\u00A0(')
                    part = part.replace(') ', ')\u00A0')
                    
                    if add_text_run_func:
                        add_text_run_func(paragraph, part, bold=bold, italic=italic)
                    else:
                        run = paragraph.add_run(part)
                        run.bold = bold
                        run.italic = italic
            else:
                formula_id = int(part)
                formula = self.latex_inlines.get(formula_id, "")
                self.add_inline(paragraph, formula, add_text_run_func, max_width=max_width, bold=bold, italic=italic)
    
    def process_block_markers(self, text: str, paragraph, add_text_run_func):
        """Process text with %%LATEX_BLOCK:id%% markers."""
        parts = []
        last_end = 0
        
        for m in self.latex.block_marker_pattern.finditer(text):
            if m.start() > last_end:
                parts.append(("text", text[last_end:m.start()]))
            parts.append(("block", int(m.group(1))))
            last_end = m.end()
        
        if last_end < len(text):
            parts.append(("text", text[last_end:]))
        
        for ptype, content in parts:
            if ptype == "text" and content.strip() and paragraph is not None:
                if self.latex.inline_marker_pattern.search(content):
                    self.process_inline_markers(content, paragraph, add_text_run_func)
                else:
                    add_text_run_func(paragraph, content)
            elif ptype == "block":
                formula = self.latex_blocks.get(content, "")
                self.add_block(formula)
    
    def cleanup(self):
        """Remove temporary files."""
        self.latex.cleanup()
