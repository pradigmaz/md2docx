"""
Main DOCX document builder.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .lists import ListBuilder
from .formulas import FormulaBuilder


class DocumentBuilder:
    """Builds DOCX document from parsed markdown."""
    
    DEFAULT_SETTINGS = {
        "fontSize": 14,
        "fontFamily": "Times New Roman",
        "lineSpacing": 1.5,
        "firstLineIndent": 1.27,
        "marginTop": 2,
        "marginBottom": 2,
        "marginLeft": 3,
        "marginRight": 1.5,
    }
    
    def __init__(self, settings: dict = None):
        self.settings = {**self.DEFAULT_SETTINGS, **(settings or {})}
        self.doc = Document()
        self._setup_document()
        
        # Initialize sub-builders
        self.lists = ListBuilder(self.doc, self.settings)
        self.formulas = FormulaBuilder(self.doc, self.settings)
    
    # Expose formula-related attributes for backward compatibility
    @property
    def latex(self):
        return self.formulas.latex
    
    @property
    def latex_blocks(self):
        return self.formulas.latex_blocks
    
    @latex_blocks.setter
    def latex_blocks(self, value):
        self.formulas.latex_blocks = value
    
    @property
    def latex_inlines(self):
        return self.formulas.latex_inlines
    
    @latex_inlines.setter
    def latex_inlines(self, value):
        self.formulas.latex_inlines = value
    
    def _setup_document(self):
        """Configure document margins and default styles."""
        section = self.doc.sections[0]
        section.top_margin = Cm(self.settings["marginTop"])
        section.bottom_margin = Cm(self.settings["marginBottom"])
        section.left_margin = Cm(self.settings["marginLeft"])
        section.right_margin = Cm(self.settings["marginRight"])
        
        style = self.doc.styles["Normal"]
        style.font.name = self.settings["fontFamily"]
        style.font.size = Pt(self.settings["fontSize"])
        
        pf = style.paragraph_format
        pf.line_spacing = self.settings["lineSpacing"]
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.first_line_indent = Cm(self.settings["firstLineIndent"])
        
        rFonts = style.element.xpath(".//w:rFonts")[0]
        rFonts.set(qn("w:ascii"), self.settings["fontFamily"])
        rFonts.set(qn("w:hAnsi"), self.settings["fontFamily"])
    
    def add_heading(self, text: str, level: int):
        """Add heading paragraph."""
        p = self.doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Cm(0)
        return p
    
    def add_paragraph(self, justify: bool = True):
        """Add regular paragraph."""
        p = self.doc.add_paragraph()
        # Use JUSTIFY alignment for proper text formatting
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.first_line_indent = Cm(self.settings["firstLineIndent"])
        p.paragraph_format.line_spacing = self.settings["lineSpacing"]
        return p
    
    def add_list_item(self, ordered: bool = False, restart: bool = False):
        """Add list item paragraph."""
        return self.lists.add_list_item(ordered=ordered, restart=restart)
    
    def add_code_block(self, code: str, language: str = ""):
        """Add code block without border."""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        
        run = p.add_run(code.strip())
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    
    def add_blockquote(self):
        """Add blockquote paragraph."""
        p = self.doc.add_paragraph(style="Quote")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        return p
    
    def add_page_break(self):
        """Add page break."""
        self.doc.add_page_break()
    
    def add_table(self, rows: int, cols: int):
        """Add table with proper column widths."""
        from docx.shared import Cm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Available width: A4 (21cm) - left margin (3cm) - right margin (1.5cm) = 16.5cm
        available_width = 16.5
        
        # Set column widths based on content type
        if cols == 3:
            widths = [Cm(1.5), Cm(9), Cm(6)]  # Total 16.5cm
        elif cols == 6:
            # For 6-column variant tables: Вариант, N, Density, Seed, ε, Примечание
            widths = [Cm(1.5), Cm(2.5), Cm(3), Cm(2), Cm(3), Cm(4.5)]  # Total 16.5cm
        else:
            # Equal widths for other tables
            col_width = Cm(available_width / cols)
            widths = [col_width] * cols
        
        # Set table width and layout explicitly
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        
        # Set table width
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(available_width * 567)))  # Convert cm to twips
        tblW.set(qn('w:type'), 'dxa')
        # Remove existing tblW if any
        for existing in tblPr.findall(qn('w:tblW')):
            tblPr.remove(existing)
        tblPr.append(tblW)
        
        # Set fixed table layout to prevent auto-resizing
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        # Remove existing tblLayout if any
        for existing in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(existing)
        tblPr.append(tblLayout)
        
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
        
        # Set column widths via tblGrid for proper fixed layout
        tblGrid = tbl.tblGrid
        if tblGrid is None:
            tblGrid = OxmlElement('w:tblGrid')
            tbl.insert(1, tblGrid)
        else:
            # Clear existing gridCol elements
            for child in list(tblGrid):
                tblGrid.remove(child)
        
        for width in widths:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), str(int(width.cm * 567)))
            tblGrid.append(gridCol)
        
        # Set column widths on cells
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if idx < len(widths):
                    cell.width = widths[idx]
                    # Also set cell width in XML for better compatibility
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(int(widths[idx].cm * 567)))
                    tcW.set(qn('w:type'), 'dxa')
                    # Remove existing tcW if any
                    for existing in tcPr.findall(qn('w:tcW')):
                        tcPr.remove(existing)
                    tcPr.append(tcW)
        
        # Disable autofit to keep fixed widths
        table.autofit = False
        
        return table
    
    def add_text_run(self, paragraph, text: str, bold: bool = False, italic: bool = False):
        """Add text run to paragraph."""
        run = paragraph.add_run(text)
        run.font.name = self.settings["fontFamily"]
        run.font.size = Pt(self.settings["fontSize"])
        run.bold = bold
        run.italic = italic
        return run
    
    def add_inline_formula(self, paragraph, formula: str, max_width: float = None):
        """Add inline LaTeX formula to paragraph."""
        return self.formulas.add_inline(paragraph, formula, self.add_text_run, max_width=max_width)
    
    def add_block_formula(self, formula: str):
        """Add block LaTeX formula as centered paragraph."""
        return self.formulas.add_block(formula)
    
    def process_inline_markers(self, text: str, paragraph, bold: bool = False, italic: bool = False, max_width: float = None):
        """Process text with %%LATEX_INLINE:id%% markers."""
        self.formulas.process_inline_markers(text, paragraph, self.add_text_run, bold=bold, italic=italic, max_width=max_width)
    
    def process_block_markers(self, text: str, paragraph):
        """Process text with %%LATEX_BLOCK:id%% markers."""
        self.formulas.process_block_markers(text, paragraph, self.add_text_run)
    
    def save(self, path: str):
        """Save document and cleanup."""
        try:
            self.doc.save(path)
        except PermissionError:
            raise PermissionError(
                f"Не удалось сохранить файл: '{path}'\n"
                "Файл открыт в другой программе (Word). Закройте его и попробуйте снова."
            )
        finally:
            self.formulas.cleanup()
