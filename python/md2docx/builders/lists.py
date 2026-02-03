"""
List building functionality for DOCX documents.
"""

from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ListBuilder:
    """Handles list creation and numbering in DOCX documents."""
    
    def __init__(self, doc, settings: dict):
        self.doc = doc
        self.settings = settings
        self._ordered_counter = 0  # Manual counter for ordered lists
        self._current_indent = 1.5  # Track current list indent
        self._setup_list_styles()
    
    def _setup_list_styles(self):
        """Configure list styles to have no spacing."""
        for style_name in ["List Number", "List Bullet"]:
            try:
                list_style = self.doc.styles[style_name]
                list_style.paragraph_format.space_after = Pt(0)
                list_style.paragraph_format.space_before = Pt(0)
            except KeyError:
                pass
    
    def add_list_item(self, ordered: bool = False, restart: bool = False, level: int = 0):
        """Add list item paragraph.
        
        Args:
            ordered: True for numbered list, False for bullet
            restart: True to restart numbering from 1
            level: Nesting level (0 = top level, 1 = first nested, etc.)
        """
        # Calculate indent based on level
        base_indent = 1.5
        level_indent = 0.75  # Additional indent per level
        left_indent = base_indent + (level * level_indent)
        self._current_indent = left_indent  # Save for continuation paragraphs
        
        if ordered:
            # Use manual numbering for ordered lists
            if restart:
                self._ordered_counter = 1
            else:
                self._ordered_counter += 1
            
            p = self.doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(left_indent)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            
            # Add number manually
            run = p.add_run(f"{self._ordered_counter}.\t")
            run.font.name = self.settings["fontFamily"]
            run.font.size = Pt(self.settings["fontSize"])
        else:
            # Use bullet style for unordered lists
            p = self.doc.add_paragraph(style="List Bullet")
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(left_indent)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
        
        return p
    
    def add_list_continuation(self, justify: bool = True):
        """Add continuation paragraph for list item (no bullet/number).
        
        Used for additional lines within a list item that should be
        aligned with the list content but without a marker.
        
        Args:
            justify: True for JUSTIFY alignment (long text), False for LEFT (short/formulas)
        """
        p = self.doc.add_paragraph()
        alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.alignment = alignment
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(self._current_indent)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        return p
