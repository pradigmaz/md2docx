"""
AST node type handlers.
"""

from docx.shared import Pt, Cm, RGBColor
import mistune


class NodeHandlers:
    """Handles specific AST node types."""
    
    def __init__(self, builder, inline_renderer):
        self.builder = builder
        self.inline = inline_renderer
    
    def handle_heading(self, node: dict):
        """Handle heading node."""
        level = node.get("attrs", {}).get("level", 1)
        p = self.builder.add_heading("", level)
        self.inline.render_children(node.get("children", []), p)
        
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    def handle_paragraph(self, node: dict):
        """Handle paragraph node."""
        children = node.get("children", [])
        
        # Filter out empty breaks at start/end
        while children and children[0].get("type") in ["softbreak", "linebreak"]:
            children = children[1:]
        while children and children[-1].get("type") in ["softbreak", "linebreak"]:
            children = children[:-1]
        
        if not children:
            return
        
        # Check if paragraph contains only block formula markers
        if len(children) == 1 and children[0].get("type") == "text":
            text = children[0].get("raw", "").strip()
            if self.builder.latex.block_marker_pattern.fullmatch(text):
                self.builder.process_block_markers(text, None)
                return
        
        # Check if paragraph contains block formula mixed with text
        if len(children) == 1 and children[0].get("type") == "text":
            text = children[0].get("raw", "")
            if self.builder.latex.block_marker_pattern.search(text):
                self._handle_mixed_block_paragraph(text)
                return
        
        # Split by softbreak/linebreak - each line becomes a new paragraph
        current_children = []
        for child in children:
            ctype = child.get("type")
            if ctype in ["softbreak", "linebreak"]:
                if current_children:
                    p = self.builder.add_paragraph()
                    for c in current_children:
                        self.inline.render_inline(c, p)
                    current_children = []
            else:
                current_children.append(child)
        
        if current_children:
            p = self.builder.add_paragraph()
            for c in current_children:
                self.inline.render_inline(c, p)
    
    def _handle_mixed_block_paragraph(self, text: str):
        """Handle paragraph with mixed text and block formulas."""
        parts = []
        last_end = 0
        
        for m in self.builder.latex.block_marker_pattern.finditer(text):
            if m.start() > last_end:
                parts.append(("text", text[last_end:m.start()]))
            parts.append(("block", int(m.group(1))))
            last_end = m.end()
        
        if last_end < len(text):
            parts.append(("text", text[last_end:]))
        
        for ptype, content in parts:
            if ptype == "text":
                content = content.strip()
                if content:
                    p = self.builder.add_paragraph()
                    if self.builder.latex.inline_marker_pattern.search(content):
                        self.builder.process_inline_markers(content, p)
                    else:
                        self.builder.add_text_run(p, content)
            elif ptype == "block":
                formula = self.builder.formulas.latex_blocks.get(content, "")
                self.builder.formulas.add_block(formula)
    
    def handle_list(self, node: dict, prev_list_type: str) -> str:
        """Handle list node. Returns current list type."""
        ordered = node.get("attrs", {}).get("ordered", False)
        list_type = "List Number" if ordered else "List Bullet"
        current_type = "ordered" if ordered else "bullet"
        
        children = node.get("children", [])
        
        for idx, item in enumerate(children):
            restart = (idx == 0 and ordered and prev_list_type != current_type)
            self.handle_list_item(item, list_type, restart)
        
        return current_type
    
    def handle_list_item(self, node: dict, list_type: str, is_first: bool = False, level: int = 0):
        """Handle list item node."""
        ordered = list_type == "List Number"
        p = self.builder.add_list_item(ordered=ordered, restart=(is_first and ordered), level=level)
        
        for child in node.get("children", []):
            ctype = child.get("type")
            if ctype in ["paragraph", "block_text"]:
                self.inline.render_children(child.get("children", []), p)
            elif ctype == "list":
                # Handle nested list
                nested_ordered = child.get("attrs", {}).get("ordered", False)
                nested_list_type = "List Number" if nested_ordered else "List Bullet"
                for idx, nested_item in enumerate(child.get("children", [])):
                    self.handle_list_item(nested_item, nested_list_type, is_first=(idx == 0), level=level + 1)
            elif ctype == "table":
                # Handle table inside list item
                self.handle_table(child)
            else:
                self.inline.render_children([child], p)
    
    def handle_code_block(self, node: dict):
        """Handle code block node."""
        code = node.get("raw", "")
        lang = node.get("attrs", {}).get("info", "")
        style = node.get("style", "")
        
        # Indented text without language is often just continuation text
        if style == "indent" and not lang:
            lines = code.strip().split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    p = self.builder.add_paragraph()
                    if self.builder.latex.block_marker_pattern.search(line):
                        self.builder.process_block_markers(line, p)
                    elif self.builder.latex.inline_marker_pattern.search(line):
                        self._parse_and_render_inline_text(line, p)
                    else:
                        self._parse_and_render_inline_text(line, p)
        else:
            self.builder.add_code_block(code, lang)
    
    def _parse_and_render_inline_text(self, text: str, paragraph):
        """Parse text for inline markdown (bold, italic) and render."""
        md = mistune.create_markdown(renderer=None)
        ast = md(text)
        if ast and ast[0].get("type") == "paragraph":
            self.inline.render_children(ast[0].get("children", []), paragraph)
        else:
            self.builder.add_text_run(paragraph, text)
    
    def handle_blockquote(self, node: dict):
        """Handle blockquote node."""
        for child in node.get("children", []):
            p = self.builder.add_blockquote()
            if child.get("type") == "paragraph":
                self.inline.render_children(child.get("children", []), p)
            else:
                self.inline.render_children([child], p)
    
    def _merge_text_nodes(self, children: list) -> list:
        """Merge consecutive text nodes into one.
        
        Mistune splits text on backslashes in tables, which breaks LaTeX formulas.
        When we see a standalone backslash node, it was originally \\\\ (double backslash).
        """
        if not children:
            return children
        
        merged = []
        current_text = ""
        
        for child in children:
            if child.get("type") == "text":
                raw = child.get("raw", "")
                # If this is a standalone backslash, it was originally \\
                # (mistune splits on \\ and keeps one \ as separate node)
                if raw == "\\":
                    current_text += "\\\\"  # Restore double backslash
                else:
                    current_text += raw
            else:
                if current_text:
                    merged.append({"type": "text", "raw": current_text})
                    current_text = ""
                merged.append(child)
        
        if current_text:
            merged.append({"type": "text", "raw": current_text})
        
        return merged
    
    def handle_table(self, node: dict):
        """Handle table node."""
        head = body = None
        for child in node.get("children", []):
            if child.get("type") == "table_head":
                head = child
            elif child.get("type") == "table_body":
                body = child
        
        if not head:
            return
        
        head_cells = head.get("children", [])
        num_cols = len(head_cells)
        body_rows = body.get("children", []) if body else []
        num_rows = 1 + len(body_rows)
        
        table = self.builder.add_table(num_rows, num_cols)
        
        # Column widths in points (1cm = 28.35pt)
        if num_cols == 3:
            col_widths_pt = [1.5 * 28.35, 9 * 28.35, 6 * 28.35]
        else:
            col_widths_pt = [16.5 / num_cols * 28.35] * num_cols
        
        # Header row
        for col, cell_node in enumerate(head_cells):
            cell = table.rows[0].cells[col]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            max_w = col_widths_pt[col] - 10 if col < len(col_widths_pt) else 150
            # Merge text nodes to fix LaTeX formulas split by mistune
            children = self._merge_text_nodes(cell_node.get("children", []))
            self.inline.render_children_in_cell(children, p, max_w)
            for run in p.runs:
                run.bold = True
        
        # Body rows
        for row_idx, row_node in enumerate(body_rows):
            for col, cell_node in enumerate(row_node.get("children", [])):
                if col < num_cols:
                    cell = table.rows[row_idx + 1].cells[col]
                    p = cell.paragraphs[0]
                    p.paragraph_format.first_line_indent = Cm(0)
                    max_w = col_widths_pt[col] - 10 if col < len(col_widths_pt) else 150
                    # Merge text nodes to fix LaTeX formulas split by mistune
                    children = self._merge_text_nodes(cell_node.get("children", []))
                    self.inline.render_children_in_cell(children, p, max_w)
