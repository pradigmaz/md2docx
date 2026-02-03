"""
Block-level node handlers for markdown AST.
"""

from docx.shared import Pt, RGBColor
import mistune


class BlockHandlerMixin:
    """Mixin for handling block-level nodes."""
    
    def handle_heading(self, node: dict):
        """Handle heading node."""
        level = node.get("attrs", {}).get("level", 1)
        p = self.builder.add_heading("", level)
        self.inline.render_children(node.get("children", []), p)
        
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
    
    def handle_paragraph(self, node: dict):
        """Handle paragraph node."""
        children = node.get("children", [])
        
        # Filter out empty breaks at start/end
        while children and children[0].get("type") in ["softbreak", "linebreak"]:
            children = children[1:]
        while children and children[-1].get("type") in ["softbreak", "linebreak"]:
            children = children[:-1]
        
        if not children:
            return
        
        # Check if paragraph contains only block formula markers
        if len(children) == 1 and children[0].get("type") == "text":
            text = children[0].get("raw", "").strip()
            if self.builder.latex.block_marker_pattern.fullmatch(text):
                self.builder.process_block_markers(text, None)
                return
        
        # Check if paragraph contains block formula mixed with text
        if len(children) == 1 and children[0].get("type") == "text":
            text = children[0].get("raw", "")
            if self.builder.latex.block_marker_pattern.search(text):
                self._handle_mixed_block_paragraph(text)
                return
        
        # Split by softbreak/linebreak - each line becomes a new paragraph
        current_children = []
        for child in children:
            ctype = child.get("type")
            if ctype in ["softbreak", "linebreak"]:
                if current_children:
                    p = self.builder.add_paragraph()
                    for c in current_children:
                        self.inline.render_inline(c, p)
                    current_children = []
            else:
                current_children.append(child)
        
        if current_children:
            p = self.builder.add_paragraph()
            for c in current_children:
                self.inline.render_inline(c, p)
    
    def _handle_mixed_block_paragraph(self, text: str):
        """Handle paragraph with mixed text and block formulas."""
        parts = []
        last_end = 0
        
        for m in self.builder.latex.block_marker_pattern.finditer(text):
            if m.start() > last_end:
                parts.append(("text", text[last_end:m.start()]))
            parts.append(("block", int(m.group(1))))
            last_end = m.end()
        
        if last_end < len(text):
            parts.append(("text", text[last_end:]))
        
        for ptype, content in parts:
            if ptype == "text":
                content = content.strip()
                if content:
                    p = self.builder.add_paragraph()
                    if self.builder.latex.inline_marker_pattern.search(content):
                        self.builder.process_inline_markers(content, p)
                    else:
                        self.builder.add_text_run(p, content)
            elif ptype == "block":
                formula = self.builder.formulas.latex_blocks.get(content, "")
                self.builder.formulas.add_block(formula)
    
    def handle_code_block(self, node: dict):
        """Handle code block node."""
        code = node.get("raw", "")
        lang = node.get("attrs", {}).get("info", "")
        style = node.get("style", "")
        
        # Indented text without language is often just continuation text
        if style == "indent" and not lang:
            lines = code.strip().split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    p = self.builder.add_paragraph()
                    if self.builder.latex.block_marker_pattern.search(line):
                        self.builder.process_block_markers(line, p)
                    elif self.builder.latex.inline_marker_pattern.search(line):
                        self._parse_and_render_inline_text(line, p)
                    else:
                        self._parse_and_render_inline_text(line, p)
        else:
            self.builder.add_code_block(code, lang)
    
    def _parse_and_render_inline_text(self, text: str, paragraph):
        """Parse text for inline markdown (bold, italic) and render."""
        md = mistune.create_markdown(renderer=None)
        ast = md(text)
        if ast and ast[0].get("type") == "paragraph":
            self.inline.render_children(ast[0].get("children", []), paragraph)
        else:
            self.builder.add_text_run(paragraph, text)
    
    def handle_blockquote(self, node: dict):
        """Handle blockquote node."""
        for child in node.get("children", []):
            p = self.builder.add_blockquote()
            if child.get("type") == "paragraph":
                self.inline.render_children(child.get("children", []), p)
            else:
                self.inline.render_children([child], p)
