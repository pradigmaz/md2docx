"""
Table node handlers for markdown AST.
"""

from docx.shared import Cm


class TableHandlerMixin:
    """Mixin for handling table nodes."""
    
    def handle_table(self, node: dict):
        """Handle table node."""
        head = body = None
        for child in node.get("children", []):
            if child.get("type") == "table_head":
                head = child
            elif child.get("type") == "table_body":
                body = child
        
        if not head:
            return
        
        head_cells = head.get("children", [])
        num_cols = len(head_cells)
        body_rows = body.get("children", []) if body else []
        num_rows = 1 + len(body_rows)
        
        table = self.builder.add_table(num_rows, num_cols)
        
        # Column widths in points (1cm = 28.35pt)
        if num_cols == 3:
            col_widths_pt = [1.5 * 28.35, 9 * 28.35, 6 * 28.35]
        else:
            col_widths_pt = [16.5 / num_cols * 28.35] * num_cols
        
        # Header row
        for col, cell_node in enumerate(head_cells):
            cell = table.rows[0].cells[col]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            max_w = col_widths_pt[col] - 10 if col < len(col_widths_pt) else 150
            children = self._merge_text_nodes(cell_node.get("children", []))
            self.inline.render_children_in_cell(children, p, max_w)
            for run in p.runs:
                run.bold = True
        
        # Body rows
        for row_idx, row_node in enumerate(body_rows):
            for col, cell_node in enumerate(row_node.get("children", [])):
                if col < num_cols:
                    cell = table.rows[row_idx + 1].cells[col]
                    p = cell.paragraphs[0]
                    p.paragraph_format.first_line_indent = Cm(0)
                    max_w = col_widths_pt[col] - 10 if col < len(col_widths_pt) else 150
                    children = self._merge_text_nodes(cell_node.get("children", []))
                    self.inline.render_children_in_cell(children, p, max_w)
    
    def _merge_text_nodes(self, children: list) -> list:
        """Merge consecutive text nodes into one.
        
        Mistune splits text on backslashes in tables, which breaks LaTeX formulas.
        """
        if not children:
            return children
        
        merged = []
        current_text = ""
        
        for child in children:
            if child.get("type") == "text":
                raw = child.get("raw", "")
                if raw == "\\":
                    current_text += "\\\\"
                else:
                    current_text += raw
            else:
                if current_text:
                    merged.append({"type": "text", "raw": current_text})
                    current_text = ""
                merged.append(child)
        
        if current_text:
            merged.append({"type": "text", "raw": current_text})
        
        return merged
