"""
Inline element rendering for markdown AST.
"""

import re
from docx.shared import Pt


class InlineRenderer:
    """Renders inline markdown elements to DOCX."""
    
    # Pattern to match **bold** text that wasn't parsed by mistune
    UNPARSED_BOLD_PATTERN = re.compile(r'\*\*([^*]+)\*\*')
    
    def __init__(self, builder):
        self.builder = builder
    
    def render_children(self, children: list, paragraph):
        """Render list of child nodes."""
        for child in children:
            self.render_inline(child, paragraph)
    
    def _render_text_with_unparsed_bold(self, text: str, paragraph, bold: bool, italic: bool):
        """Render text that may contain unparsed **bold** markers."""
        last_end = 0
        
        for match in self.UNPARSED_BOLD_PATTERN.finditer(text):
            # Add text before bold
            if match.start() > last_end:
                before_text = text[last_end:match.start()]
                if self.builder.latex.inline_marker_pattern.search(before_text):
                    self.builder.formulas.process_inline_markers(before_text, paragraph, self.builder.add_text_run, bold=bold, italic=italic)
                else:
                    self.builder.add_text_run(paragraph, before_text, bold=bold, italic=italic)
            
            # Add bold text
            bold_content = match.group(1)
            if self.builder.latex.inline_marker_pattern.search(bold_content):
                self.builder.formulas.process_inline_markers(bold_content, paragraph, self.builder.add_text_run, bold=True, italic=italic)
            else:
                self.builder.add_text_run(paragraph, bold_content, bold=True, italic=italic)
            
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            remaining = text[last_end:]
            if self.builder.latex.inline_marker_pattern.search(remaining):
                self.builder.formulas.process_inline_markers(remaining, paragraph, self.builder.add_text_run, bold=bold, italic=italic)
            else:
                self.builder.add_text_run(paragraph, remaining, bold=bold, italic=italic)
    
    def render_inline(self, node: dict, paragraph, bold: bool = False, italic: bool = False):
        """Render inline node."""
        ntype = node.get("type")
        
        if ntype == "text":
            text = node.get("raw", "")
            
            if self.builder.latex.block_marker_pattern.search(text):
                self.builder.process_block_markers(text, paragraph)
            elif '**' in text and self.UNPARSED_BOLD_PATTERN.search(text):
                # Handle unparsed **bold** with markers inside
                self._render_text_with_unparsed_bold(text, paragraph, bold, italic)
            elif self.builder.latex.inline_marker_pattern.search(text):
                self.builder.formulas.process_inline_markers(text, paragraph, self.builder.add_text_run, bold=bold, italic=italic)
            else:
                self.builder.add_text_run(paragraph, text, bold=bold, italic=italic)
        
        elif ntype == "strong":
            for c in node.get("children", []):
                self.render_inline(c, paragraph, bold=True, italic=italic)
        
        elif ntype == "emphasis":
            for c in node.get("children", []):
                self.render_inline(c, paragraph, bold=bold, italic=True)
        
        elif ntype == "codespan":
            run = paragraph.add_run(node.get("raw", ""))
            run.font.name = "Courier New"
            run.font.size = Pt(12)
        
        elif ntype == "link":
            text = "".join(
                c.get("raw", "") for c in node.get("children", []) 
                if c.get("type") == "text"
            )
            url = node.get("attrs", {}).get("url", "")
            run = paragraph.add_run(f"{text} ({url})")
            run.bold = bold
            run.italic = italic
        
        elif ntype in ["softbreak", "linebreak"]:
            paragraph.add_run("\n")
        
        elif ntype == "inline_html":
            # Handle HTML tags like <br>, <br/>, <br />
            raw = node.get("raw", "").strip().lower()
            if raw in ["<br>", "<br/>", "<br />"]:
                paragraph.add_run("\n")
    
    def render_children_in_cell(self, children: list, paragraph, max_width: float):
        """Render children in a table cell with width constraint."""
        for child in children:
            self.render_inline_in_cell(child, paragraph, max_width)
    
    def render_inline_in_cell(self, node: dict, paragraph, max_width: float, 
                               bold: bool = False, italic: bool = False):
        """Render inline node in a table cell with width constraint."""
        ntype = node.get("type")
        
        if ntype == "text":
            text = node.get("raw", "")
            
            # Check for markers first (from pre-extracted formulas)
            if self.builder.latex.block_marker_pattern.search(text):
                self.builder.process_block_markers(text, paragraph)
            elif self.builder.latex.inline_marker_pattern.search(text):
                self.builder.formulas.process_inline_markers(text, paragraph, self.builder.add_text_run, bold=bold, italic=italic, max_width=max_width)
            # Check for raw $...$ formulas (not extracted in table cells)
            elif self.builder.latex.inline_pattern.search(text):
                self._render_text_with_inline_formulas(text, paragraph, bold, italic, max_width)
            else:
                self.builder.add_text_run(paragraph, text, bold=bold, italic=italic)
        
        elif ntype == "strong":
            for c in node.get("children", []):
                self.render_inline_in_cell(c, paragraph, max_width, bold=True, italic=italic)
        
        elif ntype == "emphasis":
            for c in node.get("children", []):
                self.render_inline_in_cell(c, paragraph, max_width, bold=bold, italic=True)
        
        elif ntype == "codespan":
            run = paragraph.add_run(node.get("raw", ""))
            run.font.name = "Courier New"
            run.font.size = Pt(12)
        
        elif ntype == "link":
            text = "".join(
                c.get("raw", "") for c in node.get("children", []) 
                if c.get("type") == "text"
            )
            url = node.get("attrs", {}).get("url", "")
            run = paragraph.add_run(f"{text} ({url})")
            run.bold = bold
            run.italic = italic
        
        elif ntype in ["softbreak", "linebreak"]:
            paragraph.add_run("\n")
        
        elif ntype == "inline_html":
            # Handle HTML tags like <br>, <br/>, <br />
            raw = node.get("raw", "").strip().lower()
            if raw in ["<br>", "<br/>", "<br />"]:
                paragraph.add_run("\n")
    
    def _render_text_with_inline_formulas(self, text: str, paragraph, 
                                           bold: bool, italic: bool, max_width: float):
        """Render text containing raw $...$ inline formulas."""
        pattern = self.builder.latex.inline_pattern
        last_end = 0
        
        for match in pattern.finditer(text):
            # Add text before formula
            if match.start() > last_end:
                before_text = text[last_end:match.start()]
                self.builder.add_text_run(paragraph, before_text, bold=bold, italic=italic)
            
            # Render formula
            formula = match.group(1)
            self.builder.formulas.add_inline(paragraph, formula, self.builder.add_text_run, max_width=max_width)
            last_end = match.end()
        
        # Add remaining text
        if last_end < len(text):
            self.builder.add_text_run(paragraph, text[last_end:], bold=bold, italic=italic)
